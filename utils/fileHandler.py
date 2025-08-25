# file_handler.py
import os
from io import BytesIO
from docx import Document
from typing import List, Optional
from fastapi import UploadFile, HTTPException
import pandas as pd
import threading
import logging
import shutil
from contextlib import contextmanager
import time

logger = logging.getLogger(__name__)


class FileHandler:
    # 添加文件操作锁，保护并发文件处理
    _file_lock = threading.Lock()
    
    @classmethod
    def _sanitize_filename(cls, filename: str) -> str:
        """安全化文件名，防止路径注入"""
        if not filename:
            raise ValueError("文件名不能为空")
        
        # 移除危险字符
        dangerous_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
        safe_filename = filename
        for char in dangerous_chars:
            safe_filename = safe_filename.replace(char, '_')
        
        # 限制文件名长度
        if len(safe_filename) > 255:
            name, ext = os.path.splitext(safe_filename)
            safe_filename = name[:255-len(ext)] + ext
            
        return safe_filename
    
    @classmethod
    def _ensure_directory_exists(cls, directory: str) -> None:
        """确保目录存在，如果不存在则创建"""
        try:
            if not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)
        except OSError as e:
            logger.error(f"创建目录失败: {directory}, 错误: {e}")
            raise HTTPException(status_code=500, detail="无法创建目录")
    
    @classmethod
    @contextmanager
    def _safe_file_operation(self, file_path: str, mode: str = 'wb'):
        """安全的文件操作上下文管理器"""
        file_handle = None
        temp_path = None
        try:
            # 创建临时文件路径
            temp_path = file_path + '.tmp'
            
            # 确保目录存在
            self._ensure_directory_exists(os.path.dirname(file_path))
            
            # 打开文件
            file_handle = open(temp_path, mode)
            yield file_handle
            
            # 操作成功，将临时文件重命名为目标文件
            file_handle.close()
            file_handle = None
            
            if os.path.exists(file_path):
                os.remove(file_path)
            os.rename(temp_path, file_path)
            temp_path = None
            
        except Exception as e:
            logger.error(f"文件操作失败: {file_path}, 错误: {e}")
            # 清理临时文件
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except OSError:
                    pass
            raise
        finally:
            # 确保文件句柄被关闭
            if file_handle:
                try:
                    file_handle.close()
                except Exception:
                    pass
    
    @classmethod
    def _save_docx(cls, file: UploadFile, directory: str) -> str:
        """安全保存DOCX文件"""
        file_stream = None
        try:
            # 确保文件名安全
            safe_filename = cls._sanitize_filename(file.filename)
            if not safe_filename.lower().endswith('.docx'):
                raise ValueError("文件必须是DOCX格式")
            
            file_path = os.path.join(directory, safe_filename)
            
            # 读取文件内容
            file_content = file.file.read()
            if not file_content:
                raise ValueError("文件内容为空")
            
            # 验证文件内容
            try:
                file_stream = BytesIO(file_content)
                document = Document(file_stream)
                # 检查文档是否有效
                if not document.paragraphs and not document.tables:
                    raise ValueError("文档内容无效")
            except Exception as e:
                raise ValueError(f"无效的DOCX文件: {e}")
            
            # 使用安全文件操作保存
            with cls._safe_file_operation(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"DOCX文件保存成功: {file_path}")
            return safe_filename
            
        except Exception as e:
            logger.error(f"保存DOCX文件失败: {e}")
            raise HTTPException(status_code=400, detail=f"文件保存失败: {str(e)}")
        finally:
            if file_stream:
                try:
                    file_stream.close()
                except Exception:
                    pass
            # 重置文件指针
            try:
                file.file.seek(0)
            except Exception:
                pass
    
    @classmethod
    def _save_excel(cls, file: UploadFile, directory: str) -> str:
        """安全保存Excel文件"""
        try:
            safe_filename = cls._sanitize_filename(file.filename)
            if not safe_filename.lower().endswith(('.xlsx', '.xls')):
                raise ValueError("文件必须是Excel格式")
            
            file_path = os.path.join(directory, safe_filename)
            
            # 读取并验证文件内容
            file_content = file.file.read()
            if not file_content:
                raise ValueError("文件内容为空")
            
            # 验证Excel文件
            try:
                if safe_filename.lower().endswith('.xlsx'):
                    df = pd.read_excel(BytesIO(file_content), engine='openpyxl')
                else:
                    df = pd.read_excel(BytesIO(file_content), engine='xlrd')
                
                if df.empty:
                    raise ValueError("Excel文件内容为空")
            except Exception as e:
                raise ValueError(f"无效的Excel文件: {e}")
            
            # 使用安全文件操作保存
            with cls._safe_file_operation(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"Excel文件保存成功: {file_path}")
            return safe_filename
            
        except Exception as e:
            logger.error(f"保存Excel文件失败: {e}")
            raise HTTPException(status_code=400, detail=f"文件保存失败: {str(e)}")
        finally:
            try:
                file.file.seek(0)
            except Exception:
                pass
    
    @classmethod
    def _save_csv(cls, file: UploadFile, directory: str) -> str:
        """安全保存CSV文件"""
        try:
            safe_filename = cls._sanitize_filename(file.filename)
            if not safe_filename.lower().endswith('.csv'):
                raise ValueError("文件必须是CSV格式")
            
            file_path = os.path.join(directory, safe_filename)
            
            # 读取并验证文件内容
            file_content = file.file.read()
            if not file_content:
                raise ValueError("文件内容为空")
            
            # 验证CSV文件
            try:
                df = pd.read_csv(BytesIO(file_content))
                if df.empty:
                    raise ValueError("CSV文件内容为空")
            except Exception as e:
                raise ValueError(f"无效的CSV文件: {e}")
            
            # 使用安全文件操作保存
            with cls._safe_file_operation(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"CSV文件保存成功: {file_path}")
            return safe_filename
            
        except Exception as e:
            logger.error(f"保存CSV文件失败: {e}")
            raise HTTPException(status_code=400, detail=f"文件保存失败: {str(e)}")
        finally:
            try:
                file.file.seek(0)
            except Exception:
                pass
    
    @classmethod
    def save_file(cls, file: UploadFile, directory: str) -> str:
        """保存上传的文件"""
        start_time = time.time()
        timeout = 15  # 15秒超时
        
        try:
            logger.info(f"开始保存文件: {file.filename}")
            
            # 确保目录存在
            cls._ensure_directory_exists(directory)
            
            # 根据文件类型选择保存方法
            filename = file.filename.lower()
            logger.info(f"文件类型: {filename}")
            
            if filename.endswith('.docx'):
                result = cls._save_docx(file, directory)
            elif filename.endswith(('.xlsx', '.xls')):
                result = cls._save_excel(file, directory)
            elif filename.endswith('.csv'):
                result = cls._save_csv(file, directory)
            else:
                raise ValueError(f"不支持的文件类型: {filename}")
            
            save_time = time.time() - start_time
            logger.info(f"文件 {file.filename} 保存成功，耗时: {save_time:.2f}秒")
            return result
                    
        except TimeoutError as e:
            logger.error(f"文件保存超时: {e}")
            raise HTTPException(status_code=408, detail=f"文件保存超时: {str(e)}")
        except Exception as e:
            save_time = time.time() - start_time
            logger.error(f"文件保存失败 (耗时: {save_time:.2f}秒): {e}")
            raise HTTPException(status_code=400, detail=f"文件保存失败: {str(e)}")
    
    @classmethod
    def cleanup_directory(cls, directory: str) -> None:
        """安全清理目录内容"""
        try:
            if os.path.exists(directory):
                with cls._file_lock:
                    for item in os.listdir(directory):
                        item_path = os.path.join(directory, item)
                        try:
                            if os.path.isfile(item_path):
                                os.remove(item_path)
                            elif os.path.isdir(item_path):
                                shutil.rmtree(item_path)
                        except OSError as e:
                            logger.warning(f"清理文件失败: {item_path}, 错误: {e}")
                            continue
                logger.info(f"目录清理完成: {directory}")
        except Exception as e:
            logger.error(f"目录清理失败: {directory}, 错误: {e}")
    
    @classmethod
    def safe_delete_file(cls, file_path: str) -> bool:
        """安全删除文件"""
        try:
            if os.path.exists(file_path):
                with cls._file_lock:
                    os.remove(file_path)
                logger.info(f"文件删除成功: {file_path}")
                return True
            return False
        except OSError as e:
            logger.error(f"文件删除失败: {file_path}, 错误: {e}")
            return False

    @classmethod
    def process_files(cls, directory: str, files: List[UploadFile]) -> None:
        """处理多个上传文件"""
        start_time = time.time()
        timeout = 30  # 30秒超时
        
        try:
            logger.info(f"开始处理 {len(files)} 个文件到目录: {directory}")
            
            # 确保目录存在
            cls._ensure_directory_exists(directory)
            logger.info(f"目录准备完成: {directory}")
            
            # 清理目录中的旧文件
            logger.info("开始清理旧文件...")
            cls.cleanup_directory(directory)
            logger.info("旧文件清理完成")
            
            # 处理每个文件
            for i, file in enumerate(files):
                # 检查超时
                if time.time() - start_time > timeout:
                    raise TimeoutError(f"文件处理超时 ({timeout}秒)")
                
                if not file.filename:
                    raise ValueError("文件名不能为空")
                
                logger.info(f"处理文件 {i+1}/{len(files)}: {file.filename}")
                
                # 保存文件
                saved_filename = cls.save_file(file, directory)
                logger.info(f"文件 {file.filename} 保存成功: {saved_filename}")
            
            total_time = time.time() - start_time
            logger.info(f"成功处理 {len(files)} 个文件到目录: {directory}, 耗时: {total_time:.2f}秒")
            
        except TimeoutError as e:
            logger.error(f"文件处理超时: {e}")
            raise HTTPException(status_code=408, detail=f"文件处理超时: {str(e)}")
        except Exception as e:
            total_time = time.time() - start_time
            logger.error(f"处理文件失败 (耗时: {total_time:.2f}秒): {e}")
            raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")
