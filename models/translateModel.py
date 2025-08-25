import ctranslate2
import transformers
import pandas as pd
import threading
from openpyxl import load_workbook
from utils.config_manager import config_manager
import logging
import time
import os
from fastapi import HTTPException
from docx import Document
from tqdm import tqdm
import sys


class TranslatorSingleton:
    """翻译器单例类，支持多GPU负载均衡和任务类型感知的实例分配"""
    
    _instance = None
    _initialized = False
    _cpu_instances = []
    _cuda_instances = []
    _gpu_devices = []
    _gpu_memory_info = {}
    _text_translation_instances = []
    _file_translation_instances = []
    _task_type_distribution = {"text": 0.6, "file": 0.4}
    _lock = threading.Lock()  # 添加缺失的锁
    _tokenizers = {}  # 添加缺失的tokenizers字典
    _instance_locks = []  # 添加缺失的instance_locks列表
    
    @classmethod
    def _check_cuda_support(cls):
        """检查CUDA支持 - 只使用PyTorch检查"""
        try:
            import torch
            
            # 检查PyTorch CUDA支持
            if not torch.cuda.is_available():
                logging.info("ℹ️  PyTorch CUDA不可用，将使用CPU模式")
                return False
            
            gpu_count = torch.cuda.device_count()
            if gpu_count == 0:
                logging.info("ℹ️  PyTorch未检测到GPU设备，将使用CPU模式")
                return False
            
            logging.info(f"✅ PyTorch检测到 {gpu_count} 个GPU设备")
            
            # 检查模型路径是否存在
            test_model_path = config_manager.get_model_path(config_manager.get("SETTINGS", "SEQ_TRANSLATE_MODEL"))
            if not os.path.exists(test_model_path):
                logging.warning(f"⚠️  模型路径不存在: {test_model_path}")
                return False
            
            # 尝试创建CUDA设备上的translator来测试
            try:
                test_translator = ctranslate2.Translator(
                    test_model_path,
                    device="cuda",  # ctranslate2使用"cuda"而不是"cuda:0"
                    inter_threads=1,
                    intra_threads=1,
                    compute_type="float32"  # T4使用float32
                )
                
                # 如果成功创建，说明CUDA支持正常
                del test_translator
                logging.info("✅ ctranslate2 CUDA设备测试成功，GPU模式可用")
                return True
                
            except Exception as e:
                logging.warning(f"⚠️  ctranslate2 CUDA设备测试失败: {e}")
                logging.info("ℹ️  虽然PyTorch检测到CUDA，但ctranslate2无法使用GPU，将回退到CPU模式")
                return False
                
        except ImportError:
            logging.warning("⚠️  PyTorch未安装，无法检查CUDA支持")
            return False
        except Exception as e:
            logging.warning(f"⚠️  CUDA支持检查失败: {e}")
            return False
    
    @classmethod
    def _safe_cuda_call(cls, func, *args, default=None, **kwargs):
        """安全的CUDA调用，使用PyTorch而不是ctranslate2"""
        try:
            import torch
            
            if not torch.cuda.is_available():
                return default
            
            # 对于GPU相关的调用，我们使用PyTorch
            if func.__name__ == 'get_device_count':
                return torch.cuda.device_count()
            elif func.__name__ == 'get_device_memory_info':
                gpu_id = args[0] if args else 0
                if gpu_id < torch.cuda.device_count():
                    props = torch.cuda.get_device_properties(gpu_id)
                    # 创建一个类似ctranslate2的内存信息对象
                    class MockMemoryInfo:
                        def __init__(self, total, free):
                            self.total = total
                            self.free = free
                    
                    # 获取当前内存使用情况
                    torch.cuda.set_device(gpu_id)
                    allocated = torch.cuda.memory_allocated(gpu_id)
                    reserved = torch.cuda.memory_reserved(gpu_id)
                    free = props.total_memory - allocated
                    
                    return MockMemoryInfo(props.total_memory, free)
                else:
                    return default
            else:
                # 其他CUDA调用，返回默认值
                return default
                
        except Exception as e:
            logging.warning(f"CUDA调用失败 {func.__name__}: {e}")
            return default

    @classmethod
    def initialize_models(cls, num_cpu_models=2, num_gpu_models=4):
        """初始化模型实例，支持多GPU部署，自动回退到CPU模式"""
        # 避免长时间持有类锁
        with cls._lock:
            if cls._initialized:
                return
        
        # 在锁外进行模型初始化
        cpu_instances = []
        cuda_instances = []
        instance_locks = []
        
        # 获取可用的GPU数量和设备信息
        gpu_count = 0
        cuda_support = cls._check_cuda_support()
        
        if cuda_support:
            try:
                # 使用PyTorch获取GPU信息
                import torch
                gpu_count = torch.cuda.device_count()
                cls._gpu_devices = list(range(gpu_count))
                logging.info(f"✅ 检测到 {gpu_count} 个GPU设备: {cls._gpu_devices}")
                
                # 初始化GPU内存信息
                for gpu_id in cls._gpu_devices:
                    try:
                        props = torch.cuda.get_device_properties(gpu_id)
                        total_memory = props.total_memory
                        allocated = torch.cuda.memory_allocated(gpu_id)
                        free_memory = total_memory - allocated
                        
                        cls._gpu_memory_info[gpu_id] = {
                            'total': total_memory,
                            'free': free_memory,
                            'used': allocated
                        }
                        logging.info(f"GPU {gpu_id}: {props.name}, 总内存 {total_memory/1024**3:.2f}GB, 已用 {allocated/1024**3:.2f}GB")
                    except Exception as e:
                        logging.warning(f"无法获取GPU {gpu_id} 内存信息: {e}")
                        cls._gpu_memory_info[gpu_id] = {'total': 0, 'free': 0, 'used': 0}
            except Exception as e:
                logging.warning(f"⚠️  CUDA支持检查失败，将使用CPU模式: {e}")
                gpu_count = 0
        else:
            logging.info("ℹ️  CUDA不可用，将使用CPU模式")
            gpu_count = 0
        
        # 创建CPU实例
        for i in range(num_cpu_models):
            instance_lock = threading.Lock()
            try:
                translator = ctranslate2.Translator(
                    config_manager.get_model_path(config_manager.get("SETTINGS", "SEQ_TRANSLATE_MODEL")),
                    inter_threads=4,
                    intra_threads=1
                )
                
                cpu_instances.append({
                    "translator": translator,
                    "task_count": threading.Semaphore(10),
                    "lock": instance_lock,
                    "id": f"cpu_{i}",
                    "type": "cpu",
                    "response_times": [],
                    "last_used": time.time()
                })
                instance_locks.append(instance_lock)
                logging.info(f"✅ 成功创建CPU实例: cpu_{i}")
                
            except Exception as e:
                logging.error(f"创建CPU实例失败: {e}")
                continue
        
        # 创建GPU实例，支持多GPU负载均衡
        if gpu_count > 0 and cuda_support:
            try:
                # 计算每个GPU上应该创建的实例数量，确保均匀分布
                instances_per_gpu = num_gpu_models // gpu_count
                remaining_instances = num_gpu_models % gpu_count
                
                # 计算文本翻译和文件翻译的实例数量
                total_gpu_instances = num_gpu_models
                text_instances = int(total_gpu_instances * cls._task_type_distribution["text"])
                file_instances = total_gpu_instances - text_instances
                
                text_instance_count = 0
                file_instance_count = 0
                
                for gpu_id in cls._gpu_devices:
                    # 为每个GPU分配实例数量
                    current_gpu_instances = instances_per_gpu
                    if remaining_instances > 0:
                        current_gpu_instances += 1
                        remaining_instances -= 1
                    
                    for i in range(current_gpu_instances):
                        instance_lock = threading.Lock()
                        try:
                            # 为T4显卡优化参数 - 使用正确的ctranslate2 device参数
                            translator = ctranslate2.Translator(
                                config_manager.get_model_path(config_manager.get("SETTINGS", "SEQ_TRANSLATE_MODEL")),
                                device="cuda",  # ctranslate2使用"cuda"而不是"cuda:0"
                                inter_threads=2,
                                intra_threads=1,
                                device_index=gpu_id,
                                compute_type="float32"  # T4使用float32，float16可能不被支持
                            )
                            
                            # 根据实例数量分配任务类型
                            if text_instance_count < text_instances:
                                task_type = "text"
                                text_instance_count += 1
                                cls._text_translation_instances.append(f"cuda_{gpu_id}_{i}")
                            else:
                                task_type = "file"
                                file_instance_count += 1
                                cls._file_translation_instances.append(f"cuda_{gpu_id}_{i}")
                            
                            instance_info = {
                                "translator": translator,
                                "task_count": threading.Semaphore(config_manager.getint('GPU', 'GPU_CONCURRENT_LIMIT', 4)),
                                "lock": instance_lock,
                                "id": f"cuda_{gpu_id}_{i}",
                                "type": "cuda",
                                "gpu_id": gpu_id,
                                "task_type": task_type,  # 添加任务类型标签
                                "response_times": [],
                                "last_used": time.time(),
                                "memory_usage": 0,
                                "total_tasks": 0,
                                "successful_tasks": 0,
                                "text_tasks": 0,      # 文本翻译任务计数
                                "file_tasks": 0       # 文件翻译任务计数
                            }
                            
                            cuda_instances.append(instance_info)
                            instance_locks.append(instance_lock)
                            
                            logging.info(f"✅ 成功创建GPU实例: cuda_{gpu_id}_{i} 在GPU {gpu_id}上，任务类型: {task_type}")
                            
                        except Exception as e:
                            logging.error(f"在GPU {gpu_id}上创建实例失败: {e}")
                            continue
            except Exception as e:
                logging.error(f"创建GPU实例时出错: {e}")
                logging.info("⚠️  将回退到纯CPU模式")
        else:
            logging.info("ℹ️  使用纯CPU模式，将增加CPU实例数量")
            # 如果GPU不可用，增加CPU实例数量
            additional_cpu_instances = min(4, num_gpu_models)  # 最多增加4个CPU实例
            for i in range(additional_cpu_instances):
                instance_lock = threading.Lock()
                try:
                    translator = ctranslate2.Translator(
                        config_manager.get_model_path(config_manager.get("SETTINGS", "SEQ_TRANSLATE_MODEL")),
                        inter_threads=4,
                        intra_threads=1
                    )
                    
                    cpu_instances.append({
                        "translator": translator,
                        "task_count": threading.Semaphore(10),
                        "lock": instance_lock,
                        "id": f"cpu_extra_{i}",
                        "type": "cpu",
                        "response_times": [],
                        "last_used": time.time()
                    })
                    instance_locks.append(instance_lock)
                    logging.info(f"✅ 成功创建额外CPU实例: cpu_extra_{i}")
                    
                except Exception as e:
                    logging.error(f"创建额外CPU实例失败: {e}")
                    continue
        
        # 再次获取锁并设置初始化状态
        with cls._lock:
            if not cls._initialized:  # Double-check
                cls._cpu_instances = cpu_instances
                cls._cuda_instances = cuda_instances
                cls._instance_locks = instance_locks
                cls._initialized = True
                
                if cuda_instances:
                    logging.info(f"✅ 模型初始化完成。CPU实例: {len(cpu_instances)}, GPU实例: {len(cuda_instances)}")
                else:
                    logging.info(f"✅ 模型初始化完成（CPU模式）。CPU实例: {len(cpu_instances)}")
                    logging.info("ℹ️  系统将在CPU模式下运行，性能可能较慢但功能完整")
    
    @classmethod
    def cleanup_resources(cls):
        """清理资源"""
        with cls._lock:
            if not cls._initialized:
                return
            
            # 清理CPU实例
            for instance in cls._cpu_instances:
                try:
                    if "translator" in instance:
                        del instance["translator"]
                except Exception as e:
                    logging.warning(f"清理CPU实例时出错: {e}")
            
            # 清理GPU实例
            for instance in cls._cuda_instances:
                try:
                    if "translator" in instance:
                        del instance["translator"]
                except Exception as e:
                    logging.warning(f"清理GPU实例时出错: {e}")
            
            # 清理tokenizer缓存
            for tokenizer in cls._tokenizers.values():
                try:
                    del tokenizer
                except Exception as e:
                    logging.warning(f"清理tokenizer时出错: {e}")
            
            cls._cpu_instances.clear()
            cls._cuda_instances.clear()
            cls._tokenizers.clear()
            cls._instance_locks.clear()
            cls._initialized = False
            cls._gpu_devices.clear()
            cls._gpu_memory_info.clear()
            logging.info("✅ 资源清理完成")

    @classmethod
    def _load_tokenizer(cls, src_lang: str):
        """加载tokenizer，添加内存管理"""
        with cls._lock:
            if (src_lang, "tokenizer") not in cls._tokenizers:
                try:
                    # 检查内存使用情况，如果缓存过多则清理
                    if len(cls._tokenizers) > 10:  # 限制缓存数量
                        # 清理最旧的tokenizer
                        oldest_key = min(cls._tokenizers.keys(), key=lambda k: cls._tokenizers[k].get('last_used', 0))
                        if oldest_key in cls._tokenizers:
                            del cls._tokenizers[oldest_key]
                            logging.info(f"清理旧tokenizer缓存: {oldest_key}")
                    
                    # 获取模型名称
                    model_name = config_manager.get("SETTINGS", "SEQ_TRANSLATE_MODEL")
                    
                    # 尝试从TOKENIZER_LIST获取tokenizer路径
                    tokenizer_path = None
                    if config_manager.config.has_section('TOKENIZER_LIST'):
                        tokenizer_path = config_manager.config.get('TOKENIZER_LIST', model_name, fallback=None)
                    
                    if tokenizer_path and os.path.exists(tokenizer_path):
                        # 使用配置的tokenizer路径
                        logging.info(f"使用配置的tokenizer路径: {tokenizer_path}")
                        tokenizer = transformers.AutoTokenizer.from_pretrained(tokenizer_path)
                    else:
                        # 回退到模型路径
                        model_path = config_manager.get_model_path(model_name)
                        logging.info(f"使用模型路径加载tokenizer: {model_path}")
                        tokenizer = transformers.AutoTokenizer.from_pretrained(model_path)
                    
                    cls._tokenizers[(src_lang, "tokenizer")] = {
                        'tokenizer': tokenizer,
                        'last_used': time.time()
                    }
                    logging.info(f"加载tokenizer成功: {src_lang}")
                except Exception as e:
                    logging.error(f"加载tokenizer失败: {e}")
                    raise
            
            # 更新使用时间
            cls._tokenizers[(src_lang, "tokenizer")]['last_used'] = time.time()
            return cls._tokenizers[(src_lang, "tokenizer")]['tokenizer']

    @classmethod
    def _get_least_loaded_model(cls, use_cuda=False, task_type="text"):
        """获取负载最轻的模型实例，支持任务类型感知的负载均衡策略"""
        start_time = time.time()
        instances = cls._cuda_instances if use_cuda else cls._cpu_instances
        
        # 性能监控：记录实例选择开始
        logging.info(f"🚀 [负载均衡] 开始选择模型实例 - 任务类型: {task_type}, 使用CUDA: {use_cuda}")
        logging.info(f"📊 [负载均衡] 可用实例数量: {len(instances)}")
        
        if not instances:
            raise RuntimeError("No model instances available. Call initialize_models() first.")
        
        # 根据任务类型筛选实例
        if use_cuda and task_type:
            # 优先选择对应任务类型的实例
            preferred_instances = [inst for inst in instances if inst.get("task_type") == task_type]
            fallback_instances = [inst for inst in instances if inst.get("task_type") != task_type]
            
            # 如果首选实例可用，使用首选实例
            if preferred_instances:
                instances = preferred_instances
                logging.info(f"🎯 [负载均衡] 使用{task_type}专用实例，数量: {len(instances)}")
            else:
                # 如果首选实例不可用，使用备用实例
                instances = fallback_instances
                logging.warning(f"⚠️ [负载均衡] {task_type}专用实例不可用，使用备用实例，数量: {len(instances)}")
        
        # 使用智能负载均衡算法
        selected_instance = None
        best_score = float('-inf')
        
        # 更新GPU内存使用情况
        if use_cuda:
            cls._update_gpu_memory_usage(instances)
            logging.info(f"🔄 [GPU监控] 已更新GPU内存使用信息")
        
        # 记录所有实例的详细状态
        logging.info(f"📋 [负载均衡] 开始评估实例状态...")
        for instance in instances:
            try:
                # 使用非阻塞方式检查信号量可用性
                if instance["task_count"].acquire(blocking=False):
                    # 立即释放，只是检查可用性
                    instance["task_count"].release()
                    
                    # 计算综合评分（越高越好）
                    score = 0
                    score_details = []
                    
                    # 1. 可用槽位评分（权重40%）
                    available_slots = instance["task_count"]._value
                    max_slots = config_manager.getint('GPU', 'GPU_CONCURRENT_LIMIT', 4)
                    slot_score = (available_slots / max_slots) * 40
                    score += slot_score
                    score_details.append(f"槽位: {available_slots}/{max_slots} ({slot_score:.1f})")
                    
                    # 2. 任务类型匹配奖励（权重20%）
                    if use_cuda and "task_type" in instance:
                        if instance["task_type"] == task_type:
                            score += 20  # 任务类型匹配奖励
                            score_details.append(f"类型匹配: +20")
                        else:
                            score -= 10  # 任务类型不匹配惩罚
                            score_details.append(f"类型不匹配: -10")
                    
                    # 3. 响应时间评分（权重25%）
                    if instance["response_times"]:
                        avg_response_time = sum(instance["response_times"]) / len(instance["response_times"])
                        # 响应时间越短，评分越高
                        response_score = max(0, (1000 - avg_response_time) / 1000 * 25)
                        score += response_score
                        score_details.append(f"响应时间: {avg_response_time:.1f}ms ({response_score:.1f})")
                    
                    # 4. 任务成功率评分（权重10%）
                    if "total_tasks" in instance and instance["total_tasks"] > 0:
                        success_rate = instance["successful_tasks"] / instance["total_tasks"]
                        success_score = success_rate * 10
                        score += success_score
                        score_details.append(f"成功率: {success_rate:.2f} ({success_score:.1f})")
                    
                    # 5. GPU内存使用评分（权重5%）
                    if use_cuda and "gpu_id" in instance:
                        gpu_id = instance["gpu_id"]
                        if gpu_id in cls._gpu_memory_info:
                            memory_usage_ratio = cls._gpu_memory_info[gpu_id]['used'] / cls._gpu_memory_info[gpu_id]['total']
                            # 内存使用率越低，评分越高
                            memory_score = (1 - memory_usage_ratio) * 5
                            score += memory_score
                            score_details.append(f"GPU内存: {memory_usage_ratio:.2f} ({memory_score:.1f})")
                    
                    # 记录实例详细评分
                    logging.info(f"📊 [负载均衡] 实例 {instance['id']} 评分: {score:.1f} - {' | '.join(score_details)}")
                    
                    # 选择评分最高的实例
                    if score > best_score:
                        best_score = score
                        selected_instance = instance
                        logging.info(f"🏆 [负载均衡] 新最佳实例: {instance['id']} (评分: {score:.1f})")
                        
            except Exception as e:
                logging.warning(f"❌ [负载均衡] 检查实例 {instance.get('id', 'unknown')} 负载时出错: {e}")
                continue
        
        # 如果没有找到可用实例，使用轮询策略
        if selected_instance is None:
            # 随机选择一个实例并等待
            import random
            selected_instance = random.choice(instances)
            logging.warning(f"🔄 [负载均衡] 使用轮询策略选择实例: {selected_instance['id']}")
        
        # 记录最终选择的实例
        selection_time = (time.time() - start_time) * 1000
        logging.info(f"✅ [负载均衡] 实例选择完成 - 选择: {selected_instance['id']}, 耗时: {selection_time:.1f}ms")
        
        # 尝试获取选中的实例
        try:
            # 设置超时时间，避免无限等待
            if selected_instance["task_count"].acquire(blocking=True, timeout=30):
                selected_instance["last_used"] = time.time()
                if "total_tasks" in selected_instance:
                    selected_instance["total_tasks"] += 1
                
                # 记录任务类型统计
                if use_cuda and "task_type" in selected_instance:
                    if task_type == "text":
                        selected_instance["text_tasks"] += 1
                    elif task_type == "file":
                        selected_instance["file_tasks"] += 1
                
                return selected_instance
            else:
                raise RuntimeError("获取模型实例超时")
        except Exception as e:
            logging.error(f"获取模型实例失败: {e}")
            # 如果获取失败，尝试其他实例
            for instance in instances:
                if instance != selected_instance:
                    try:
                        if instance["task_count"].acquire(blocking=False):
                            instance["last_used"] = time.time()
                            if "total_tasks" in instance:
                                instance["total_tasks"] += 1
                            
                            # 记录任务类型统计
                            if use_cuda and "task_type" in instance:
                                if task_type == "text":
                                    instance["text_tasks"] += 1
                                elif task_type == "file":
                                    instance["file_tasks"] += 1
                            
                            return instance
                    except:
                        continue
            raise RuntimeError("无法获取任何可用的模型实例")

    @classmethod
    def _update_gpu_memory_usage(cls, gpu_instances):
        """更新GPU内存使用情况"""
        try:
            import torch
            
            for instance in gpu_instances:
                if "gpu_id" in instance:
                    gpu_id = instance["gpu_id"]
                    try:
                        # 使用torch获取GPU内存信息
                        if torch.cuda.is_available() and gpu_id < torch.cuda.device_count():
                            torch.cuda.set_device(gpu_id)
                            memory_info = torch.cuda.get_device_properties(gpu_id)
                            total_memory = memory_info.total_memory
                            allocated_memory = torch.cuda.memory_allocated(gpu_id)
                            cached_memory = torch.cuda.memory_reserved(gpu_id)
                            free_memory = total_memory - allocated_memory
                            
                            cls._gpu_memory_info[gpu_id] = {
                                'total': total_memory,
                                'free': free_memory,
                                'used': allocated_memory,
                                'cached': cached_memory
                            }
                            instance["memory_usage"] = allocated_memory
                        else:
                            logging.debug(f"GPU {gpu_id} 不可用或超出范围")
                    except Exception as e:
                        logging.warning(f"更新GPU {gpu_id} 内存信息失败: {e}")
        except Exception as e:
            logging.warning(f"更新GPU内存使用情况失败: {e}")

    @classmethod
    def _release_model(cls, model_instance):
        """释放模型实例"""
        try:
            if model_instance and "task_count" in model_instance:
                model_instance["task_count"].release()
                
                # 记录响应时间
                if "last_used" in model_instance:
                    response_time = time.time() - model_instance["last_used"]
                    model_instance["response_times"].append(response_time)
                    
                    # 只保留最近的10个响应时间
                    if len(model_instance["response_times"]) > 10:
                        model_instance["response_times"] = model_instance["response_times"][-10:]
                    
                    # 更新最后使用时间
                    model_instance["last_used"] = time.time()
                    
                    # 记录任务成功
                    if "successful_tasks" in model_instance:
                        model_instance["successful_tasks"] += 1
                    
        except Exception as e:
            logging.error(f"释放模型实例时出错: {e}")

    @classmethod
    def translate_sentence(cls, text: str, src_lang: str, tgt_lang: str, use_cuda=False, via_eng=False, task_type="text") -> str:
        """翻译单个句子"""
        model_instance = None
        start_time = time.time()
        
        # 性能监控：记录翻译开始
        logging.info(f"🚀 [翻译] 开始翻译单个句子 - 任务类型: {task_type}, 使用CUDA: {use_cuda}")
        logging.info(f"📝 [翻译] 源语言: {src_lang} -> 目标语言: {tgt_lang}, 文本长度: {len(text)}")
        
        try:
            if via_eng and src_lang != "eng_Latn" and tgt_lang != "eng_Latn":
                # 防止无限递归：检查语言组合
                if src_lang == tgt_lang:
                    raise ValueError("源语言和目标语言不能相同")
                
                logging.info(f"🔄 [翻译] 通过英语中转翻译: {src_lang} -> eng_Latn -> {tgt_lang}")
                # First translate to English
                intermediate_text = cls.translate_sentence(text, src_lang, "eng_Latn", use_cuda, False, task_type)
                # Then translate from English to target language
                return cls.translate_sentence(intermediate_text, "eng_Latn", tgt_lang, use_cuda, False, task_type)

            # 获取模型实例
            instance_start_time = time.time()
            model_instance = cls._get_least_loaded_model(use_cuda, task_type)
            instance_time = (time.time() - instance_start_time) * 1000
            logging.info(f"✅ [翻译] 获取模型实例完成 - 实例ID: {model_instance['id']}, 耗时: {instance_time:.1f}ms")
            
            translator = model_instance["translator"]
            tokenizer = cls._load_tokenizer(src_lang)

            # 使用安全的tokenization方法
            token_start_time = time.time()
            source = cls._safe_encode_and_convert(tokenizer, text)
            token_time = (time.time() - token_start_time) * 1000
            logging.debug(f"🔤 [翻译] Tokenization结果: {source[:10]}... (共{len(source)}个), 耗时: {token_time:.1f}ms")
            
            # 映射语言代码到NLLB格式
            mapped_tgt_lang = cls._map_language_code(tgt_lang)
            target_prefix = [[mapped_tgt_lang]]  # 恢复：需要是List[List[str]]格式，每个文本对应一个列表
            
            # 执行翻译
            inference_start_time = time.time()
            results = translator.translate_batch([source], target_prefix=[target_prefix])
            inference_time = (time.time() - inference_start_time) * 1000
            logging.info(f"⚡ [翻译] 模型推理完成 - 耗时: {inference_time:.1f}ms")
            
            # 修复decode调用 - 根据官方demo处理结果
            try:
                if hasattr(results[0], 'hypotheses') and results[0].hypotheses:
                    # 根据官方demo：跳过第一个token
                    target = results[0].hypotheses[0][1:]
                    
                    # 使用convert_tokens_to_ids + decode
                    if hasattr(tokenizer, 'convert_tokens_to_ids') and hasattr(tokenizer, 'decode'):
                        target_ids = tokenizer.convert_tokens_to_ids(target)
                        translated_text = tokenizer.decode(target_ids)
                    else:
                        # 回退：直接使用hypotheses
                        translated_text = ' '.join(target)
                else:
                    # 回退处理
                    translated_text = str(results[0])
                    
            except Exception as e:
                logging.error(f"处理翻译结果失败: {e}")
                # 回退到原始结果
                translated_text = str(results[0])
            
            # 记录响应时间
            response_time = time.time() - start_time
            if "response_times" in model_instance:
                model_instance["response_times"].append(response_time)
                # 只保留最近的10个响应时间
                if len(model_instance["response_times"]) > 10:
                    model_instance["response_times"] = model_instance["response_times"][-10:]
            
            return translated_text
            
        except Exception as e:
            # 记录错误
            logging.error(f"翻译句子时发生错误: {e}")
            raise HTTPException(status_code=500, detail=f"翻译服务内部错误: {e}")
        finally:
            if model_instance:
                cls._release_model(model_instance)

    @classmethod
    def translate_batch(cls, texts: list, src_lang: str, tgt_lang: str, use_cuda=False, via_eng=False, task_type="text") -> list:
        """批量翻译"""
        model_instance = None
        start_time = time.time()
        
        # 性能监控：记录批量翻译开始
        logging.info(f"🚀 [批量翻译] 开始批量翻译 - 任务类型: {task_type}, 使用CUDA: {use_cuda}")
        logging.info(f"📝 [批量翻译] 源语言: {src_lang} -> 目标语言: {tgt_lang}, 文本数量: {len(texts)}")
        logging.info(f"📊 [批量翻译] 总文本长度: {sum(len(str(t)) for t in texts)} 字符")
        
        try:
            if via_eng and src_lang != "eng_Latn" and tgt_lang != "eng_Latn":
                # 防止无限递归：检查语言组合
                if src_lang == tgt_lang:
                    raise ValueError("源语言和目标语言不能相同")
                
                logging.info(f"🔄 [批量翻译] 通过英语中转翻译: {src_lang} -> eng_Latn -> {tgt_lang}")
                # First translate to English
                intermediate_texts = cls.translate_batch(texts, src_lang, "eng_Latn", use_cuda, False, task_type)
                # Then translate from English to target language
                return cls.translate_batch(intermediate_texts, "eng_Latn", tgt_lang, use_cuda, False, task_type)

            # 获取模型实例
            instance_start_time = time.time()
            model_instance = cls._get_least_loaded_model(use_cuda, task_type)
            instance_time = (time.time() - instance_start_time) * 1000
            logging.info(f"✅ [批量翻译] 获取模型实例完成 - 实例ID: {model_instance['id']}, 耗时: {instance_time:.1f}ms")
            translator = model_instance["translator"]
            tokenizer = cls._load_tokenizer(src_lang)

            # 使用安全的tokenization方法
            sources = []
            for i, text in enumerate(texts):
                source = cls._safe_encode_and_convert(tokenizer, text)
                logging.debug(f"文本{i}: Tokenization结果: {source[:10]}... (共{len(source)}个)")
                sources.append(source)
            
            # 修复target_prefix格式 - 根据官方demo，每个文本都需要对应的目标语言前缀
            # ctranslate2期望: List[Optional[List[str]]] 格式
            mapped_tgt_lang = cls._map_language_code(tgt_lang)  # 映射语言代码
            target_prefix = [[mapped_tgt_lang] for _ in texts]  # 恢复：批量翻译需要这个格式
            
            # 执行批量翻译 - 修复API调用格式
            # ctranslate2期望: List[List[str]] 格式的tokenized文本
            results = translator.translate_batch(sources, target_prefix=target_prefix)
            
            # 修复decode调用 - 根据官方demo处理结果
            translated_texts = []
            for hyp in results:
                try:
                    if hasattr(hyp, 'hypotheses') and hyp.hypotheses:
                        # 根据官方demo：跳过第一个token
                        target = hyp.hypotheses[0][1:]
                        
                        # 使用convert_tokens_to_ids + decode
                        if hasattr(tokenizer, 'convert_tokens_to_ids') and hasattr(tokenizer, 'decode'):
                            target_ids = tokenizer.convert_tokens_to_ids(target)
                            translated_text = tokenizer.decode(target_ids)
                        else:
                            # 回退：直接使用hypotheses
                            translated_text = ' '.join(target)
                    else:
                        # 回退处理
                        translated_text = str(hyp)
                    
                    translated_texts.append(translated_text)
                    
                except Exception as e:
                    logging.error(f"处理翻译结果失败: {e}")
                    # 回退到原始结果
                    translated_texts.append(str(hyp))
            
            # 记录响应时间
            response_time = time.time() - start_time
            if "response_times" in model_instance:
                model_instance["response_times"].append(response_time)
                # 只保留最近的10个响应时间
                if len(model_instance["response_times"]) > 10:
                    model_instance["response_times"] = model_instance["response_times"][-10:]
            
            return translated_texts
            
        except Exception as e:
            # 记录错误
            logging.error(f"批量翻译时发生错误: {e}")
            raise HTTPException(status_code=500, detail=f"批量翻译服务内部错误: {e}")
        finally:
            if model_instance:
                cls._release_model(model_instance)

    @classmethod
    def get_gpu_status(cls):
        """获取GPU状态信息"""
        try:
            # 更新GPU内存使用情况
            cls._update_gpu_memory_usage(cls._cuda_instances)
            
            status = {
                "gpu_count": len(cls._gpu_devices),
                "gpu_devices": cls._gpu_devices,
                "gpu_instances": len(cls._cuda_instances),
                "cpu_instances": len(cls._cpu_instances),
                "memory_info": cls._gpu_memory_info,
                "instance_details": [],
                "performance_metrics": {},
                "task_type_distribution": cls._task_type_distribution,
                "text_instances": cls._text_translation_instances,
                "file_instances": cls._file_translation_instances,
                "timestamp": time.time()
            }
            
            # 添加实例详细信息
            total_tasks = 0
            total_successful = 0
            total_text_tasks = 0
            total_file_tasks = 0
            avg_response_time = 0
            response_times = []
            
            for instance in cls._cuda_instances:
                instance_info = {
                    "id": instance["id"],
                    "gpu_id": instance.get("gpu_id", "N/A"),
                    "type": instance["type"],
                    "task_type": instance.get("task_type", "unknown"),
                    "available_slots": instance["task_count"]._value,
                    "response_times": instance.get("response_times", []),
                    "memory_usage": instance.get("memory_usage", 0),
                    "total_tasks": instance.get("total_tasks", 0),
                    "successful_tasks": instance.get("successful_tasks", 0),
                    "text_tasks": instance.get("text_tasks", 0),
                    "file_tasks": instance.get("file_tasks", 0),
                    "success_rate": 0
                }
                
                # 计算成功率
                if instance_info["total_tasks"] > 0:
                    instance_info["success_rate"] = instance_info["successful_tasks"] / instance_info["total_tasks"]
                
                # 计算平均响应时间
                if instance_info["response_times"]:
                    instance_avg_response = sum(instance_info["response_times"]) / len(instance_info["response_times"])
                    instance_info["avg_response_time"] = instance_avg_response
                    response_times.extend(instance_info["response_times"])
                
                status["instance_details"].append(instance_info)
                
                # 累计统计
                total_tasks += instance_info["total_tasks"]
                total_successful += instance_info["successful_tasks"]
                total_text_tasks += instance_info["text_tasks"]
                total_file_tasks += instance_info["file_tasks"]
            
            # 计算整体性能指标
            if response_times:
                avg_response_time = sum(response_times) / len(response_times)
            
            status["performance_metrics"] = {
                "total_tasks": total_tasks,
                "total_successful": total_successful,
                "total_text_tasks": total_text_tasks,
                "total_file_tasks": total_file_tasks,
                "overall_success_rate": total_successful / total_tasks if total_tasks > 0 else 0,
                "avg_response_time": avg_response_time,
                "total_instances": len(cls._cuda_instances) + len(cls._cpu_instances),
                "active_instances": sum(1 for inst in cls._cuda_instances if inst["task_count"]._value < config_manager.getint('GPU', 'GPU_CONCURRENT_LIMIT', 4)),
                "task_type_balance": {
                    "text_tasks_percentage": (total_text_tasks / total_tasks * 100) if total_tasks > 0 else 0,
                    "file_tasks_percentage": (total_file_tasks / total_tasks * 100) if total_tasks > 0 else 0
                }
            }
            
            return status
        except Exception as e:
            logging.error(f"获取GPU状态失败: {e}")
            return {"error": str(e)}

    @classmethod
    def _map_language_code(cls, lang_code: str) -> str:
        """映射语言代码到NLLB模型期望的格式"""
        language_mapping = {
            # 前端语言代码 -> NLLB官方语言代码
            "zh_Hans": "zho_Hans",  # 中文简体 -> NLLB官方代码
            "zh_Hant": "zho_Hant",  # 中文繁体 -> NLLB官方代码
            "en": "eng_Latn",       # 英语
            "eng": "eng_Latn",      # 英语
            "eng_Latn": "eng_Latn", # 英语（前端格式）
            "mon": "khk_Cyrl",      # 蒙古语 -> NLLB官方代码
            "mn": "khk_Cyrl",       # 蒙古语 -> NLLB官方代码
            "khk_Cyrl": "khk_Cyrl", # 西里尔蒙文（前端格式）-> NLLB官方代码
            "khk": "khk_Cyrl",      # 蒙古语简写 -> NLLB官方代码
            # 如果已经是NLLB官方格式，直接返回
            "zho_Hans": "zho_Hans", # NLLB中文简体
            "zho_Hant": "zho_Hant", # NLLB中文繁体
            "eng_Latn": "eng_Latn", # NLLB英语
            "khk_Cyrl": "khk_Cyrl", # NLLB蒙古语
        }
        return language_mapping.get(lang_code, lang_code)

    @classmethod
    def _simple_tokenize(cls, text: str) -> list:
        """简单的文本分割，绕过tokenizer问题"""
        try:
            # 尝试使用transformers tokenizer
            if 'transformers' in sys.modules:
                try:
                    from transformers import AutoTokenizer
                    # 使用默认的tokenizer
                    tokenizer = AutoTokenizer.from_pretrained("facebook/nllb-200-distilled-600M")
                    tokens = tokenizer.tokenize(text)
                    if tokens:
                        return tokens
                except Exception as e:
                    logging.warning(f"Transformers tokenizer失败: {e}")
            
            # 回退到简单分割
            return text.split()
            
        except Exception as e:
            logging.warning(f"简单tokenize失败: {e}")
            return text.split()
    
    @classmethod
    def _safe_encode_and_convert(cls, tokenizer, text: str) -> list:
        """安全的编码和转换，处理各种异常情况"""
        try:
            # 方法1：尝试标准流程
            try:
                raw_tokens = tokenizer.encode(text)
                logging.debug(f"标准encode成功: 类型={type(raw_tokens)}, 值={raw_tokens}")
                
                # 确保是整数列表
                if isinstance(raw_tokens, (list, tuple)):
                    token_ids = []
                    for item in raw_tokens:
                        try:
                            token_ids.append(int(item))
                        except (ValueError, TypeError):
                            logging.warning(f"跳过无效token: {item}")
                    
                    if token_ids:
                        tokens = tokenizer.convert_ids_to_tokens(token_ids)
                        if tokens:
                            return tokens
                
            except Exception as e:
                logging.warning(f"标准流程失败: {e}")
            
            # 方法2：尝试直接tokenize
            try:
                if hasattr(tokenizer, 'tokenize'):
                    tokens = tokenizer.tokenize(text)
                    if tokens:
                        return tokens
            except Exception as e:
                logging.warning(f"直接tokenize失败: {e}")
            
            # 方法3：使用简单分割
            logging.info("使用简单文本分割作为回退")
            return text.split()
            
        except Exception as e:
            logging.error(f"所有tokenization方法都失败: {e}")
            return text.split()


class DocxTranslator(TranslatorSingleton):
    @staticmethod
    def translate_run(run, src_lang, tgt_lang, via_eng=False):
        if not run.text.strip():
            return ""

        translated_lines = TranslatorSingleton.translate_batch(texts=run.text.split("\n"),
                                                              src_lang=src_lang,
                                                              tgt_lang=tgt_lang,
                                                              use_cuda=True,
                                                              via_eng=via_eng,
                                                              task_type="file")  # 使用文件翻译实例
        return '\n'.join(translated_lines)

    @staticmethod
    def translate_paragraph(paragraph, src_lang, tgt_lang, via_eng=False):
        translated_runs = []

        for run in paragraph.runs:
            translated_text = DocxTranslator.translate_run(run=run,
                                                           src_lang=src_lang,
                                                           tgt_lang=tgt_lang,
                                                           via_eng=via_eng)
            translated_runs.append((translated_text, run))

        paragraph.clear()

        for translated_text, original_run in translated_runs:
            translated_run = paragraph.add_run(translated_text)

            translated_run.bold = original_run.bold
            translated_run.italic = original_run.italic
            translated_run.underline = original_run.underline
            translated_run.font.size = original_run.font.size
            translated_run.font.name = original_run.font.name
            translated_run.font.color.rgb = original_run.font.color.rgb
            translated_run.font.highlight_color = original_run.font.highlight_color

    @staticmethod
    def translate_docx(input_path: str, output_path: str, src_lang: str, tgt_lang: str, via_eng=False):
        doc = Document(input_path)
        translated_doc = Document()

        for para in tqdm(doc.paragraphs, desc=f"Translating {input_path}"):
            if para.text.strip():
                DocxTranslator.translate_paragraph(paragraph=para,
                                                   src_lang=src_lang,
                                                   tgt_lang=tgt_lang,
                                                   via_eng=via_eng)
                translated_doc.add_paragraph(para.text)

        translated_doc.save(output_path)


class TableTranslator(TranslatorSingleton):
    @staticmethod
    def translate_text(text, src_lang, tgt_lang, via_eng=False):
        if text is None:
            return text
        lines = text.split('\n')
        translated_lines = TranslatorSingleton.translate_batch(texts=lines,
                                                               src_lang=src_lang,
                                                               tgt_lang=tgt_lang,
                                                               use_cuda=True,
                                                               via_eng=via_eng,
                                                               task_type="file")  # 使用文件翻译实例
        return '\n'.join(translated_lines)

    @staticmethod
    def translate_excel(input_path: str, output_path: str, src_lang: str, tgt_lang: str, via_eng=False):
        """
        优化后的Excel翻译功能
        支持所有工作表的翻译，正确处理合并单元格和只读单元格
        """
        try:
            print(f"开始翻译Excel文件: {input_path}")
            wb = load_workbook(input_path)
            
            # 获取工作表信息
            total_sheets = len(wb.worksheets)
            print(f"Excel文件包含 {total_sheets} 个工作表")
            
            # 统计所有工作表的翻译信息
            total_cells_translated = 0
            total_texts_translated = 0
            
            for sheet_idx, sheet in enumerate(wb.worksheets, 1):
                print(f"\n{'='*50}")
                print(f"正在翻译工作表 {sheet_idx}/{total_sheets}: {sheet.title}")
                print(f"工作表尺寸: {sheet.max_row} 行 x {sheet.max_column} 列")
                
                # 存储需要翻译的单元格内容
                cells_to_translate = []
                
                # 第一步：收集所有需要翻译的文本内容
                print(f"正在扫描工作表 {sheet.title} 的单元格...")
                for row in sheet.iter_rows():
                    for cell in row:
                        try:
                            # 安全地获取单元格值
                            cell_value = cell.value
                            
                            # 只翻译非空字符串
                            if cell_value and isinstance(cell_value, str) and cell_value.strip():
                                # 检查是否为合并单元格的主单元格
                                is_merge_master = False
                                for merged_range in sheet.merged_cells.ranges:
                                    if cell.coordinate == merged_range.start_cell.coordinate:
                                        is_merge_master = True
                                        break
                                
                                # 只翻译合并单元格的主单元格或非合并单元格
                                if is_merge_master or not any(cell.coordinate in merged_range for merged_range in sheet.merged_cells.ranges):
                                    cells_to_translate.append({
                                        'coordinate': cell.coordinate,
                                        'value': cell_value,
                                        'is_merged': is_merge_master
                                    })
                        except Exception as e:
                            print(f"警告：无法读取单元格 {cell.coordinate}: {e}")
                            continue
                
                print(f"工作表 {sheet.title} 找到 {len(cells_to_translate)} 个需要翻译的单元格")
                total_cells_translated += len(cells_to_translate)
                
                # 第二步：批量翻译文本
                if cells_to_translate:
                    # 提取所有文本进行批量翻译
                    texts_to_translate = [cell['value'] for cell in cells_to_translate]
                    print(f"开始批量翻译 {len(texts_to_translate)} 个文本...")
                    
                    try:
                        translated_texts = TranslatorSingleton.translate_batch(
                            texts=texts_to_translate,
                            src_lang=src_lang,
                            tgt_lang=tgt_lang,
                            use_cuda=True,
                            via_eng=via_eng,
                            task_type="file"
                        )
                        
                        # 第三步：将翻译结果应用到工作表
                        print(f"正在将翻译结果应用到工作表 {sheet.title}...")
                        for i, cell_info in enumerate(cells_to_translate):
                            try:
                                coord = cell_info['coordinate']
                                translated_text = translated_texts[i] if i < len(translated_texts) else cell_info['value']
                                
                                # 设置单元格值
                                sheet[coord].value = translated_text
                                
                                # 如果是合并单元格，需要应用到整个合并范围
                                if cell_info['is_merged']:
                                    for merged_range in sheet.merged_cells.ranges:
                                        if coord == merged_range.start_cell.coordinate:
                                            # 合并单元格的值会自动应用到整个范围
                                            print(f"  ✓ 已翻译合并单元格 {coord}: {cell_info['value']} -> {translated_text}")
                                            break
                                else:
                                    print(f"  ✓ 已翻译单元格 {coord}: {cell_info['value']} -> {translated_text}")
                                    
                                total_texts_translated += 1
                                    
                            except Exception as e:
                                print(f"警告：无法设置单元格 {coord} 的值: {e}")
                                continue
                    
                    except Exception as e:
                        print(f"工作表 {sheet.title} 翻译失败: {e}")
                        # 如果翻译失败，保持原值
                        continue
                else:
                    print(f"工作表 {sheet.title} 没有需要翻译的文本")
                
                # 第四步：保存工作表
                print(f"✓ 工作表 {sheet.title} 翻译完成")
            
            # 保存工作簿
            print(f"\n{'='*50}")
            print(f"所有工作表翻译完成，正在保存文件...")
            wb.save(output_path)
            print(f"✓ Excel文件翻译完成，已保存到: {output_path}")
            print(f"📊 翻译统计:")
            print(f"  - 总工作表数: {total_sheets}")
            print(f"  - 总单元格数: {total_cells_translated}")
            print(f"  - 总翻译文本: {total_texts_translated}")
            
        except Exception as e:
            print(f"Excel翻译过程中出现错误: {e}")
            import traceback
            traceback.print_exc()
            raise e

    @staticmethod
    def translate_csv(input_path: str, output_path: str, src_lang: str, tgt_lang: str, via_eng=False):
        """
        优化后的CSV翻译功能
        支持大文件处理和错误恢复
        """
        try:
            print(f"开始翻译CSV文件: {input_path}")
            
            # 读取CSV文件
            df = pd.read_csv(input_path)
            print(f"CSV文件包含 {len(df)} 行，{len(df.columns)} 列")
            print(f"列名: {list(df.columns)}")
            
            # 收集所有需要翻译的文本
            texts_to_translate = []
            text_positions = []  # 记录文本在DataFrame中的位置
            
            print("正在扫描CSV文件中的文本内容...")
            for col_idx, column in enumerate(df.columns):
                print(f"  扫描列 {col_idx+1}/{len(df.columns)}: {column}")
                for row_idx, value in enumerate(df[column]):
                    if value is not None and isinstance(value, str) and value.strip():
                        texts_to_translate.append(value)
                        text_positions.append((row_idx, col_idx))
            
            print(f"找到 {len(texts_to_translate)} 个需要翻译的文本")
            
            if texts_to_translate:
                try:
                    print(f"开始批量翻译 {len(texts_to_translate)} 个文本...")
                    # 批量翻译
                    translated_texts = TranslatorSingleton.translate_batch(
                        texts=texts_to_translate,
                        src_lang=src_lang,
                        tgt_lang=tgt_lang,
                        use_cuda=True,
                        via_eng=via_eng,
                        task_type="file"
                    )
                    
                    print(f"翻译完成，正在将结果应用到CSV文件...")
                    # 将翻译结果应用到DataFrame
                    for i, (row_idx, col_idx) in enumerate(text_positions):
                        if i < len(translated_texts):
                            df.iloc[row_idx, col_idx] = translated_texts[i]
                            print(f"  ✓ 已翻译: {texts_to_translate[i]} -> {translated_texts[i]}")
                        else:
                            print(f"警告：翻译结果不足，保持原值: {texts_to_translate[i]}")
                    
                except Exception as e:
                    print(f"翻译过程中出现错误: {e}")
                    # 如果翻译失败，保持原值
                    pass
            else:
                print("CSV文件中没有需要翻译的文本")
            
            # 保存翻译后的CSV文件
            print("正在保存翻译后的CSV文件...")
            df.to_csv(output_path, index=False)
            print(f"✓ CSV文件翻译完成，已保存到: {output_path}")
            print(f"📊 翻译统计:")
            print(f"  - 总行数: {len(df)}")
            print(f"  - 总列数: {len(df.columns)}")
            print(f"  - 翻译文本数: {len(texts_to_translate)}")
            
        except Exception as e:
            print(f"CSV翻译过程中出现错误: {e}")
            import traceback
            traceback.print_exc()
            raise e
